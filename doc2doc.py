import numpy as np
from docx import Document
from tqdm.autonotebook import tqdm
from text_similarity import text_similarity

def doc2doc(path1, path2):
    document1 = Document(open(path1, "rb")).paragraphs
    table1 = Document(open(path1, "rb")).tables
    document2 = Document(open(path2, "rb")).paragraphs
    table2 = Document(open(path2, "rb")).tables

    table_text1 = []
    for i in range(len(table1)):
        for j in range(len(table1[i].rows)):
            for k in range(len(table1[i].rows[j].cells)):
                # if len(today_file[i].rows[j].cells[k].text) > 50:
                table_text1.append(table1[i].rows[j].cells[k].text.replace("\n", " ")) 

    table_text2 = []
    for i in range(len(table2)):
        for j in range(len(table2[i].rows)):
            for k in range(len(table2[i].rows[j].cells)):
                # if len(today_file[i].rows[j].cells[k].text) > 50:
                table_text2.append(table2[i].rows[j].cells[k].text.replace("\n", " "))

    new_table1 = list(set(table_text1))
    new_table2 = list(set(table_text2))
    
    ready_today1 = []
    for i in range(len(document1)):
        # if len(document1[i].text) > 49:
        ready_today1.append(document1[i].text)

    ready_today2 = []
    for i in range(len(document2)):
        # if len(document2[i].text) > 49:
        ready_today2.append(document2[i].text)

    last1 = ready_today1 + new_table1
    last2 = ready_today2 + new_table2

    row = len(last1)
    column = len(last2)
    matrix = np.zeros((row, column))

    for i in tqdm(range(row)):
        for j in range(column):
            matrix[i, j] = text_similarity(last1[i], last2[j])
    
    mat = []
    for i in range(matrix.shape[0]):
        if matrix[i].max() != 1:
            mat.append(i)
    
    result = []
    for p in mat:
        result.append(last1[p])
    
    return result